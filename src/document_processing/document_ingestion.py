"""
Document Ingestion Module

Processes HR policy documents (Word .docx files) using:
1. Azure Document Intelligence (OCR-enabled, for scanned docs)
2. Azure Content Understanding (advanced semantic extraction)
3. Local python-docx fallback (no Azure required)
"""

import logging
import os
import re
import hashlib
import subprocess
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# Optional Azure imports
try:
    from azure.ai.documentintelligence import DocumentIntelligenceClient
    from azure.identity import AzureCliCredential, ChainedTokenCredential, ManagedIdentityCredential
    DOCINT_AVAILABLE = True
except ImportError:
    DOCINT_AVAILABLE = False
    logger.info("azure-ai-documentintelligence not installed, using local processing only")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed, Word document processing limited")


class DocumentIngestionAgent:
    """
    Processes HR policy documents (Word .docx files) and extracts text content.

    Supports:
    - Azure Document Intelligence for OCR-enabled extraction
    - Local python-docx for standard Word files
    """

    def __init__(
        self,
        use_azure: bool = False,
    ):
        self.use_azure = use_azure
        self.azure_endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")

        if self.use_azure and not self.azure_endpoint:
            logger.warning("Azure Document Intelligence not configured. Falling back to local processing.")
            self.use_azure = False

    def process_document(self, file_path: str) -> dict[str, Any]:
        """
        Process a single document and extract text content.

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with extracted text, metadata, and page info
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        ext = path.suffix.lower()
        logger.info(f"Processing document: {path.name} (type: {ext})")

        if self.use_azure and DOCINT_AVAILABLE and ext in (".docx", ".pdf"):
            return self._process_with_azure(file_path)
        elif ext == ".doc":
            # Old binary .doc format: Azure DI doesn't support it properly,
            # so use antiword directly for clean text extraction
            return self._process_doc_with_antiword(file_path)
        elif ext == ".docx" and DOCX_AVAILABLE:
            return self._process_docx_locally(file_path)
        elif ext == ".txt":
            return self._process_text_file(file_path)
        else:
            logger.warning(f"Unsupported file type: {ext}. Attempting text extraction.")
            return self._process_text_file(file_path)

    def _process_with_azure(self, file_path: str) -> dict[str, Any]:
        """Process document using Azure Document Intelligence."""
        try:
            if not DOCINT_AVAILABLE:
                raise ImportError("Azure Document Intelligence dependencies not available")
            from azure.ai.documentintelligence import DocumentIntelligenceClient as DIClient
            from azure.identity import AzureCliCredential, ChainedTokenCredential, ManagedIdentityCredential
            assert self.azure_endpoint is not None
            credential = ChainedTokenCredential(
                ManagedIdentityCredential(),
                AzureCliCredential(),
            )
            client = DIClient(
                endpoint=self.azure_endpoint,
                credential=credential,
            )

            with open(file_path, "rb") as f:
                poller = client.begin_analyze_document(
                    model_id="prebuilt-layout",
                    body=f,
                    content_type="application/octet-stream",
                )
                result = poller.result()

            text = result.content if result.content else ""
            tables = []
            if result.tables:
                for table in result.tables:
                    tables.append({
                        "row_count": table.row_count,
                        "column_count": table.column_count,
                    })

            return {
                "text": text,
                "page_count": len(result.pages) if result.pages else 1,
                "word_count": len(text.split()),
                "char_count": len(text),
                "tables": tables,
                "extraction_method": "azure_document_intelligence",
            }

        except Exception as e:
            logger.warning(f"Azure processing failed: {e}. Falling back to local processing.")
            if Path(file_path).suffix.lower() == ".docx" and DOCX_AVAILABLE:
                return self._process_docx_locally(file_path)
            return self._process_text_file(file_path)

    def _process_docx_locally(self, file_path: str) -> dict[str, Any]:
        """Process a Word document using python-docx."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for .docx processing")
        import docx as _docx
        doc = _docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        # Extract tables
        tables = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            tables.append("\n".join(table_text))

        if tables:
            text += "\n\n--- Tables ---\n\n" + "\n\n".join(tables)

        return {
            "text": text,
            "page_count": 1,
            "word_count": len(text.split()),
            "char_count": len(text),
            "tables": [{"content": t} for t in tables],
            "extraction_method": "python_docx",
        }

    def _process_doc_with_antiword(self, file_path: str) -> dict[str, Any]:
        """Process an old binary .doc file using antiword."""
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout
                return {
                    "text": text,
                    "page_count": 1,
                    "word_count": len(text.split()),
                    "char_count": len(text),
                    "tables": [],
                    "extraction_method": "antiword",
                }
            logger.warning(f"antiword returned no output for {file_path}: {result.stderr}")
        except FileNotFoundError:
            logger.warning("antiword not installed. Install with: sudo apt-get install antiword")
        except Exception as e:
            logger.warning(f"antiword failed for {file_path}: {e}")
        # Last resort — will produce garbled output for binary .doc files
        return self._process_text_file(file_path)

    def _process_text_file(self, file_path: str) -> dict[str, Any]:
        """Process a plain text file."""
        with open(file_path, "r", errors="replace") as f:
            text = f.read()

        return {
            "text": text,
            "page_count": 1,
            "word_count": len(text.split()),
            "char_count": len(text),
            "tables": [],
            "extraction_method": "text_read",
        }


def generate_document_id(file_path: str) -> str:
    """Generate a deterministic document ID from file path."""
    return hashlib.md5(file_path.encode()).hexdigest()


def extract_policy_number(filename: str) -> str:
    """Extract policy number from filename (e.g., '50410' from '50410 - Hiring...')."""
    match = re.match(r"^(\d+)", filename)
    return match.group(1) if match else ""


def categorize_policy(filename: str) -> str:
    """Categorize policy based on filename keywords."""
    lower = filename.lower()
    if "hiring" in lower or "probation" in lower or "rehiring" in lower:
        return "hiring"
    if "leave" in lower or "pto" in lower or "disability" in lower:
        return "leave"
    if "career" in lower:
        return "career_path"
    if "hours" in lower or "pay" in lower or "holiday" in lower:
        return "compensation"
    if "uniform" in lower or "dress" in lower or "operational" in lower:
        return "operational"
    if "ethics" in lower or "code of" in lower:
        return "ethics"
    if "blood" in lower or "pathogen" in lower or "safety" in lower:
        return "safety"
    return "general"
